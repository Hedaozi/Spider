# -*- coding: utf-8 -*-
"""
Created on Tue Nov 19 20:37:17 2019

@author: 何凌锋
"""

import tkinter as tk
import tkinter.messagebox
from selenium import webdriver
import time
import datetime
import locale
import random
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn
import operator
import pandas
import numpy

def test():
    tkinter.messagebox.showinfo(title = 'Test', message = 'Test')
    return
    
def start_chorme():
    global browser
    try:
        browser = webdriver.Chrome()
        browser.get("https://s.weibo.com/")
    except:
        tkinter.messagebox.showerror(title = '错误', message = '请安装Google Chrome和对应版本的Webdriver，并将浏览器所在文件夹加入系统环境变量。或者检查网络连接。')
    return

def gen_res():
    res = dict({
        'fulltext':'', 'time':'', 'reed':0, 'comed':0, 'liked':0,
        're_flag':False, 're_text':'', 're_reed':0, 're_comed':0, 're_liked':0
    })
    return res

def get_fulltext(target_text):
    if "feed_list_content_full" in target_text:
        p_temp = """feed_list_content_full.+?收起全文"""
    else:
        p_temp = """feed_list_content.+?</p>"""
    pattern_temp = re.compile(p_temp)
    text_temp = pattern_temp.findall(target_text)
    if text_temp == []:
        return "未找到正文"
    text_temp = re.sub("<.+?>", "", text_temp[0])
    text_temp = re.sub("feed.+?>", "", text_temp)    
    text_temp = re.sub("<ahref.+?target=", "", text_temp)
    text_temp = text_temp.replace("收起全文", "").replace("微博内容", "").replace("转发微博","")
    return text_temp

def get_repeat(target_text):
    p_temp = "转发\d+"
    pattern_temp = re.compile(p_temp)
    text_temp = pattern_temp.findall(target_text)
    if text_temp == []:
        return 0
    else:
        text_temp = text_temp[0].replace("转发", "")
        return int(text_temp)

def get_comment(target_text):
    p_temp = "评论\d+"
    pattern_temp = re.compile(p_temp)
    text_temp = pattern_temp.findall(target_text)
    if text_temp == []:
        return 0
    else:
        text_temp = text_temp[0].replace("评论", "")
        return int(text_temp)

def get_like(target_text):
    p_temp = """<iclass="icon-acticon-act-praise"></i><em>\d+?</em></a></li>"""
    pattern_temp = re.compile(p_temp)
    text_temp = pattern_temp.findall(target_text)
    if text_temp == []:
        return 0
    else:
        text_temp = text_temp[0].replace("""<iclass="icon-acticon-act-praise"></i><em>""", "").replace("</em></a></li>", "")
        return int(text_temp)

def get_time(target_text):
    time_now = time.localtime(time.time())
    todaydate = str(time_now.tm_year) + "年" + str(time_now.tm_mon) + "月" + str(time_now.tm_mday) + "日"
    time_yest = datetime.date.today() + datetime.timedelta(-1)
    yestdate = str(time_yest.year) + "年" + str(time_yest.month) + "月" + str(time_yest.day) + "日"
    p_temp = "\d+?年\d+?月\d+日\d+?:\d+"
    pattern_temp = re.compile(p_temp)
    text_temp = pattern_temp.findall(target_text)
    if not (text_temp == []):
        text_temp = text_temp[0]
    else:
        p_temp = "\d+?月\d+日\d+?:\d+"
        pattern_temp = re.compile(p_temp)
        text_temp = pattern_temp.findall(target_text)
        if not (text_temp == []):
            text_temp = str(time_now.tm_year) + "年" + text_temp[0]
        else:
            p_temp = "今天\d+?:\d+"
            pattern_temp = re.compile(p_temp)
            text_temp = pattern_temp.findall(target_text)
            if not (text_temp == []):
                text_temp = todaydate + text_temp[0][2:]
            else:
                p_temp = "昨天\d+?:\d+"
                pattern_temp = re.compile(p_temp)
                text_temp = pattern_temp.findall(target_text)
                if not (text_temp == []):
                    text_temp = yestdate + text_temp[0][2:]
                else:
                    p_temp = ">\d+?小时前</a>"
                    pattern_temp = re.compile(p_temp)
                    text_temp = pattern_temp.findall(target_text)
                    if not (text_temp == []):
                        time_b = int(text_temp[0][1:-7])
                        text_temp = (datetime.datetime.now() + datetime.timedelta(hours = 0 - time_b)).strftime("%Y年%m月%d日%H:%M")
                    else:
                        p_temp = ">\d+?分钟前</a>"
                        pattern_temp = re.compile(p_temp)
                        text_temp = pattern_temp.findall(target_text)
                        if not (text_temp == []):
                            time_b = int(text_temp[0][1:-7])
                            text_temp = (datetime.datetime.now() + datetime.timedelta(minutes = 0 - time_b)).strftime("%Y年%m月%d日%H:%M")
                        else:
                            text_temp = datetime.datetime.now().strftime("%Y年%m月%d日%H:%M")
    return text_temp

def findtext(text):
    p_div = """card-wrap[\d\D]+?/card-wrap"""
    pattern_div = re.compile(p_div)
    matcher_div = pattern_div.findall(text)
    res_list = []
    for item in matcher_div:
        res_temp = gen_res()
        if "转发微博" in item:
            res_temp['re_flag'] = True
            p_temp = """微博内容[\d\D]+?<!--转发微博-->"""
            pattern_temp = re.compile(p_temp)
            text_temp = pattern_temp.findall(item)[0]
            res_temp['fulltext'] = get_fulltext(text_temp)
            p_temp = """<!--转发微博-->[\d\D]+?<!--/转发微博-->"""
            pattern_temp = re.compile(p_temp)
            text_temp = pattern_temp.findall(item)[0]
            res_temp['re_text'] = get_fulltext(text_temp)
            res_temp['re_reed'] = get_repeat(text_temp)
            res_temp['re_comed'] = get_comment(text_temp)
            res_temp['re_liked'] = get_like(text_temp)
            res_temp['re_time'] = get_time(text_temp)
            p_temp = """<!--/转发微博-->[\d\D]+?/card-wrap"""
            pattern_temp = re.compile(p_temp)
            text_temp = pattern_temp.findall(item)[0]
            res_temp['reed'] = get_repeat(text_temp)
            res_temp['comed'] = get_comment(text_temp)
            res_temp['liked'] = get_like(text_temp)
            res_temp['time'] = get_time(text_temp)
        else:
            res_temp['fulltext'] = get_fulltext(item)
            res_temp['reed'] = get_repeat(item)
            res_temp['comed'] = get_comment(item)
            res_temp['liked'] = get_like(item)
            res_temp['time'] = get_time(item)
        res_list.append(res_temp)
    return res_list
    
def start_spider():
    global browser
    global text
    global res
    text = ""
    kw = keyword.get()
    page_e = pageend.get()
    if page_e == "":
        tkinter.messagebox.showwarning(title = '警告', message = '默认终止页码为50。')
    try:
        page_e = int(page_e)
    except:
        tkinter.messagebox.showerror(title = '错误', message = '请在终止页码输入框中输入1~50的整数。')
        return
    if page_e > 50:
        page_e = 50
        tkinter.messagebox.showwarning(title = '警告', message = '微博返回结果最多显示50页。已自动将终止页吗设置为50。')
    if page_e < 1:
        tkinter.messagebox.showerror(title = '错误', message = '请在终止页码输入框中输入1~50的整数。')
        return
    if browser == 0:
        tkinter.messagebox.showerror(title = '错误', message = '请点击上侧按钮启动浏览器，并手动完成登录操作。请保持该浏览器。')
        return
    try:
        browser.get("https://s.weibo.com/")
        browser.find_element_by_xpath('/html/body/div[1]/div[2]/div/div[2]/div/input').send_keys(kw)
        browser.find_element_by_xpath('/html/body/div[1]/div[2]/div/div[2]/button').click()
        currenturl = browser.current_url[:-11] + "nodup=1&page="
        theme = "(" + kw + ")"
        page_s = 1
        for i in range(page_s, page_e + 1):
            browser.get(currenturl + str(i))
            text = text + browser.page_source + '\n'
            time.sleep(random.randrange(3,7))
        tkinter.messagebox.showinfo(title = '反馈', message = '爬取完毕，可以解析并导出。')
    except:
        tkinter.messagebox.showerror(title = '错误', message = 'Something Wrong.')
        return
    res = findtext(text.replace('\n', '').replace(' ', ''))
    return
    
def readlocal():
    global text
    global res
    road = readroad.get()
    text = open(road, "r", encoding = "utf-8").read()
    tkinter.messagebox.showinfo(title = '反馈', message = '读取成功。' + road)
    res = findtext(text.replace('\n', '').replace(' ', ''))
    return

def add_table(item, document):
    table = document.add_table(rows = 4, cols = 4, style = 'Table Grid')
    table.cell(0, 0).merge(table.cell(0, 3))
    table.cell(3, 0).merge(table.cell(3, 3))
    table.cell(0, 0).paragraphs[0].add_run("微博内容：")
    table.cell(0, 0).add_paragraph(item['fulltext'])
    table.cell(1, 0).paragraphs[0].add_run("时间")
    table.cell(1, 1).paragraphs[0].add_run(str(item['time']))
    table.cell(1, 2).paragraphs[0].add_run("转发数")
    table.cell(1, 3).paragraphs[0].add_run(str(item['reed']))
    table.cell(2, 0).paragraphs[0].add_run("评论数")
    table.cell(2, 1).paragraphs[0].add_run(str(item['comed']))
    table.cell(2, 2).paragraphs[0].add_run("点赞数")
    table.cell(2, 3).paragraphs[0].add_run(str(item['liked']))
    if not (item['re_flag']):
        table.cell(3, 0).paragraphs[0].add_run("转发的微博：无")
        return
    else:
        table.add_row()
        table.add_row()
        table.cell(3, 0).paragraphs[0].add_run("转发的微博：")
        table.cell(3, 0).add_paragraph(item['re_text'])
        table.cell(4, 0).paragraphs[0].add_run("时间")
        table.cell(4, 1).paragraphs[0].add_run(str(item['re_time']))
        table.cell(4, 2).paragraphs[0].add_run("转发数")
        table.cell(4, 3).paragraphs[0].add_run(str(item['re_reed']))
        table.cell(5, 0).paragraphs[0].add_run("评论数")
        table.cell(5, 1).paragraphs[0].add_run(str(item['re_comed']))
        table.cell(5, 2).paragraphs[0].add_run("点赞数")
        table.cell(5, 3).paragraphs[0].add_run(str(item['re_liked']))
        return

def Export_1():
    global text
    kw = keyword.get()
    f = open("source results of (" + kw + ").txt", "w", encoding = "utf-8")
    f.write(text)
    f.close()
    tkinter.messagebox.showinfo(title = '反馈', message = '成功导出：\n' + "source results of (" + kw + ").txt")
    return

def Export_2():
    global text
    kw = keyword.get()
    f = open("source results of (" + kw + ") without split.txt", "w", encoding = "utf-8")
    f.write(text.replace(' ', '').replace('\n', ''))
    f.close()
    tkinter.messagebox.showinfo(title = '反馈', message = '成功导出：\n' + "source results of (" + kw + ") without split.txt")
    return

def Export_3():
    global res
    kw = keyword.get()
    f = open("result of (" + kw + ").txt", "w", encoding = "utf-8")
    i = 1
    for item in res:
        output_temp = "编号：" + str(i) + "\n"
        output_temp = output_temp + "微博正文：\n" + item['fulltext'] + "\n"
        output_temp = output_temp + "时间：" + item['time'] + " "
        output_temp = output_temp + "转发数：" + str(item['reed']) + " "
        output_temp = output_temp + "评论数：" + str(item['comed']) + " "
        output_temp = output_temp + "点赞数：" + str(item['liked']) + "\n\n"
        output_temp = output_temp + "是否转发了微博："
        if item['re_flag']:
            output_temp = output_temp + "是；被转发的微博为：\n" + item['re_text'] + "\n"
            output_temp = output_temp + "时间：" + item['re_time'] + " "
            output_temp = output_temp + "转发数：" + str(item['re_reed']) + " "
            output_temp = output_temp + "评论数：" + str(item['re_comed']) + " "
            output_temp = output_temp + "点赞数：" + str(item['re_liked'])
        else:
            output_temp = output_temp + "否"
        output_temp = output_temp + "\n\n\n\n"
        f.write(output_temp)
        i = i + 1
    f.close()
    tkinter.messagebox.showinfo(title = '反馈', message = '成功导出：\n' + "result of (" + kw + ").txt")
    return

def Export_4():
    global text
    global res
    kw = keyword.get()
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.name = u'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')
    for i in range(0, len(res)):
        p = document.add_paragraph()
        p.add_run('No. ' + str(i + 1)).bold = True
        add_table(res[i], document)
        document.add_paragraph()
    document.save("result of (" + kw + ").docx")
    tkinter.messagebox.showinfo(title = '反馈', message = '成功导出：\n' + "result of (" + kw + ").docx")
    return
    
def Export_result():
    global text
    global res
    kw = keyword.get()
    f1 = var1.get()
    f2 = var2.get()
    f3 = var3.get()
    f4 = var4.get()
    if not f1 and not f2 and not f3 and not f4:
        tkinter.messagebox.showinfo(title = '反馈', message = '无文件导出。')
        return
    infotext = "成功导出：\n"
    if f1:
        f = open("source results of (" + kw + ").txt", "w", encoding = "utf-8")
        f.write(text)
        f.close()
        infotext = infotext + "source results of (" + kw + ").txt\n"
    if f2:
        f = open("source results of (" + kw + ") without split.txt", "w", encoding = "utf-8")
        f.write(text.replace(' ', '').replace('\n', ''))
        f.close()
        infotext = infotext + "source results of (" + kw + ") without split.txt\n"
    if f3:
        f = open("result of (" + kw + ").txt", "w", encoding = "utf-8")
        i = 1
        for item in res:
            output_temp = "编号：" + str(i) + "\n"
            output_temp = output_temp + "微博正文：\n" + item['fulltext'] + "\n"
            output_temp = output_temp + "时间：" + item['time'] + " "
            output_temp = output_temp + "转发数：" + str(item['reed']) + " "
            output_temp = output_temp + "评论数：" + str(item['comed']) + " "
            output_temp = output_temp + "点赞数：" + str(item['liked']) + "\n\n"
            output_temp = output_temp + "是否转发了微博："
            if item['re_flag']:
                output_temp = output_temp + "是；被转发的微博为：\n" + item['re_text'] + "\n"
                output_temp = output_temp + "时间：" + item['re_time'] + " "
                output_temp = output_temp + "转发数：" + str(item['re_reed']) + " "
                output_temp = output_temp + "评论数：" + str(item['re_comed']) + " "
                output_temp = output_temp + "点赞数：" + str(item['re_liked'])
            else:
                output_temp = output_temp + "否"
            output_temp = output_temp + "\n\n\n\n"
            f.write(output_temp)
            i = i + 1
        f.close()
        infotext = infotext + "result of (" + kw + ").txt\n"
    if f4:
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.name = u'Times New Roman'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')
        for i in range(0, len(res)):
            p = document.add_paragraph()
            p.add_run('No. ' + str(i + 1)).bold = True
            add_table(res[i], document)
            document.add_paragraph()
        document.save("result of (" + kw + ").docx")
        infotext = infotext + "result of (" + kw + ").docx"
    tkinter.messagebox.showinfo(title = '反馈', message = infotext)
    return

def transtime(timeod):
    p = '\d+?年'
    pattern = re.compile(p)
    year = pattern.findall(timeod)[0][:-1]
    p = '\d+?月'
    pattern = re.compile(p)
    mon = pattern.findall(timeod)[0][:-1]
    if (int(mon) < 10):
        mon = "0" + mon
    p = '\d+?日'
    pattern = re.compile(p)
    day = pattern.findall(timeod)[0][:-1]
    if (int(day) < 10):
        day = "0" + day
    timeed = year + mon + day + " " + timeod[-5:]
    return timeed

def Export_csv():
    global res
    res2 = []
    for item in res:
        res2_temp = dict({'text': '', 'time': ''})
        p = '.+?//@'
        pattern = re.compile(p)
        if len(pattern.findall(item['fulltext'])) > 0:
            res2_temp['text'] = pattern.findall(item['fulltext'])[0][:-3]
        else:
            res2_temp['text'] = item['fulltext']
        res2_temp['time'] = transtime(item['time'])
        res2.append(res2_temp)
    res2 = sorted(res2, key = operator.itemgetter('time'))
    res3 = pandas.DataFrame(columns = ['time', 'text'])
    for i in range(0, len(res2)):
        res3.loc[i] = {'text': res2[i]['text'], 'time': res2[i]['time']}
    res3.to_csv("result of (" + keyword.get() + ").csv", index = False)
    return
   
browser = 0
text = ""
res = []
locale.setlocale(locale.LC_CTYPE, 'chinese')

window = tk.Tk()
window.title('微博爬虫')
window.geometry('570x380')

menubar = tk.Menu(window)
savemenu = tk.Menu(menubar, tearoff = 0)
startmenu = tk.Menu(menubar, tearoff = 0)
menubar.add_cascade(label = '开始', menu = startmenu)
startmenu.add_command(label = '启动浏览器', command = start_chorme)
startmenu.add_command(label = '开始抓取', command = start_spider)
startmenu.add_separator()
startmenu.add_command(label = '读取本地源代码', command = readlocal)
menubar.add_cascade(label = '导出', menu = savemenu)
savesources = tk.Menu(savemenu)
savemenu.add_cascade(label = '源代码', menu = savesources, underline = 0)
savesources.add_command(label = '源码', command = Export_1)
savesources.add_command(label = '''源码（去除了 \' \' and '\\n\'）''', command = Export_2)
savemenu.add_command(label = '*.txt', command = Export_3)
savemenu.add_command(label = '*.docx', command = Export_4)
menubar.add_cascade(label = '帮助', command = test)
window.config(menu = menubar)

l_webspider = tk.Label(window, text = '网络抓取源代码', fg = 'black', font = ('宋体', 15), width = 14, height = 1)
l_webspider.place(x = 10, y = 10)
b_start = tk.Button(window, text = '启动浏览器', font = ('宋体', 15), fg = 'black', width = 10, height = 1, command = start_chorme)
b_start.place(x = 10, y = 45)
l_start = tk.Label(window, text = '（请手动完成登陆操作）', fg = 'black', font = ('宋体', 15), width = 22, height = 1)
l_start.place(x = 120, y = 50)
l_keyword = tk.Label(window, text = '关键词  ', fg = 'black', font = ('宋体', 15), width = 8, height = 1)
l_pageend = tk.Label(window, text = '终止页码', fg = 'black', font = ('宋体', 15), width = 8, height = 1)
l_keyword.place(x = 10, y = 90)
l_pageend.place(x = 10, y = 130)
keyword = tk.Entry(window, show = None, fg = 'black', font = ('宋体', 15), width = 30)
pageend = tk.Entry(window, show = None, fg = 'black', font = ('宋体', 15), width = 30)
keyword.place(x = 110, y = 90)
pageend.place(x = 110, y = 130)
b_webspider = tk.Button(window, text = '开始抓取', font = ('宋体', 15), fg = 'black', width = 10, height = 1, command = start_spider)
b_webspider.place(x = 440, y = 105)

l_localread = tk.Label(window, text = '本地读取源代码', fg = 'black', font = ('宋体', 15), width = 14, height = 1)
l_localread.place(x = 10, y = 170)
l_readroad = tk.Label(window, text = '文件路径', fg = 'black', font = ('宋体', 15), width = 8, height = 1)
l_readroad.place(x = 10, y = 210)
readroad = tk.Entry(window, show = None, fg = 'black', font = ('宋体', 15), width = 30)
readroad.place(x = 110, y = 210)
b_webspider = tk.Button(window, text = '读取', fg = 'black', font = ('宋体', 15), width = 10, height = 1, command = readlocal)
b_webspider.place(x = 440, y = 205)

l_output = tk.Label(window, text = '导出选项', fg = 'black', font = ('宋体', 15), width = 8, height = 1)
l_output.place(x = 10, y = 250)
var1 = tk.IntVar()
var2 = tk.IntVar()
var3 = tk.IntVar()
var4 = tk.IntVar()
c1 = tk.Checkbutton(window, text = '源码', fg = 'black', variable = var1, font = ('宋体', 15), width = 5, height = 1, onvalue = 1, offvalue = 0)
c1.place(x = 5, y = 290)
c2 = tk.Checkbutton(window, text = '源码（去除空格和换行）', fg = 'black', variable = var2, font = ('宋体', 15), width = 23, height = 1, onvalue = 1, offvalue = 0)
c2.place(x = 5, y = 330)
c3 = tk.Checkbutton(window, text = '*.txt', fg = 'black', variable = var3, font = ('宋体', 15), width = 6, height = 1, onvalue = 1, offvalue = 0)
c3.place(x = 255, y = 290)
c4 = tk.Checkbutton(window, text = '*.docx', fg = 'black', variable = var4, font = ('宋体', 15), width = 7, height = 1, onvalue = 1, offvalue = 0)
c4.place(x = 255, y = 330)
b_output = tk.Button(window, text = '导出', fg = 'black', font = ('宋体', 15), width = 10, height = 1, command = Export_result)
b_output.place(x = 440, y = 290)
b_output = tk.Button(window, text = 'csv', fg = 'black', font = ('宋体', 15), width = 10, height = 1, command = Export_csv)
b_output.place(x = 440, y = 330)

window.mainloop()